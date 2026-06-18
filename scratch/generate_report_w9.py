import os
import shutil
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_run_font(run, name="Times New Roman", size_pt=13, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, size_pt=14, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, size_pt=13, bold=True)
    return p

def add_paragraph_styled(doc, text, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         space_before=0, space_after=6, line_spacing=1.5, bullet=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    if bullet:
        p.paragraph_format.left_indent = Inches(0.25)
        run_bullet = p.add_run("• ")
        set_run_font(run_bullet, bold=True)
        
    if bold_prefix:
        run_prefix = p.add_run(bold_prefix)
        set_run_font(run_prefix, bold=True)
        
    run_text = p.add_run(text)
    set_run_font(run_text)
    return p

def add_image_styled(doc, img_path, caption):
    if not os.path.exists(img_path):
        print(f"Warning: Image {img_path} not found!")
        return None
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture(img_path, width=Inches(4.5))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(8)
    
    run_cap = p_cap.add_run(caption)
    set_run_font(run_cap, size_pt=11, italic=True)
    return p_img

def main():
    src_dir = "/Users/dangminhtam/Documents/DoAn"
    docx_src = os.path.join(src_dir, "Báo cáo tuần/23730156-23730071-Trần Văn Chót-Dương Thị Quỳnh Châm-Báo cáo tuần 8.docx")
    docx_dst = os.path.join(src_dir, "Báo cáo tuần/23730156-23730071-Trần Văn Chót-Dương Thị Quỳnh Châm-Báo cáo tuần 9.docx")
    pic_dir = os.path.join(src_dir, "Picture")
    
    print(f"Copying {docx_src} to {docx_dst}...")
    shutil.copyfile(docx_src, docx_dst)
    
    doc = docx.Document(docx_dst)
    
    # 1. Update Title
    doc.paragraphs[0].text = "" # clear
    run_title = doc.paragraphs[0].add_run("BÁO CÁO TUẦN 9 ĐỒ ÁN")
    set_run_font(run_title, size_pt=16, bold=True)
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Delete all paragraphs after the table (keep Table 0 and first 3 paragraphs)
    while len(doc.paragraphs) > 3:
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
        
    # Add empty paragraph space after Table 0
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(6)
    p_space.paragraph_format.space_after = Pt(6)
    
    # 3. Add Week 9 Content
    
    # --- SECTION 1 ---
    add_heading_styled(doc, "1. Nội dung công việc trong tuần", level=1)
    
    add_paragraph_styled(doc, 
        "Lập trình máy chủ Python Flask hoạt động tại cổng 8088 của PC/Mac trung tâm, thiết lập cơ sở dữ liệu quan hệ SQLite (tệp faces.db) để quản lý đồng bộ thông tin ID - Name và xây dựng các webhook nhận ảnh chân dung gốc JPEG cùng lịch sử log điểm danh từ thiết bị nhúng gửi lên.",
        bold_prefix="• Xây dựng hệ thống máy chủ trung tâm (Flask Server + SQLite DB): "
    )
    
    add_paragraph_styled(doc, 
        "Thiết lập cơ chế truyền thông điệp UDP socket (cổng 8090) cho phép thiết bị biên ESP32-S3 tự động gửi gói tin quảng bá (Broadcast) để tìm kiếm IP của máy chủ trung tâm trong mạng nội bộ LAN, loại bỏ hoàn toàn việc phải cấu hình IP thủ công vào firmware trước khi nạp.",
        bold_prefix="• Phát triển giao thức kết nối mạng tự động UDP Auto-Discovery: "
    )
    
    add_paragraph_styled(doc, 
        "Hỗ trợ đồng bộ hóa danh sách người dùng hai chiều giữa thiết bị nhúng và máy chủ. Lập trình kịch bản đăng ký khuôn mặt nhanh trực tiếp trên thiết bị (Double click nút BOOT/GPIO 0) hoặc thông qua giao diện Web Dashboard, tự động chụp ảnh và truyền file chân dung gốc về server.",
        bold_prefix="• Triển khai tính năng Đăng ký nhanh và Đồng bộ hóa dữ liệu: "
    )
    
    add_paragraph_styled(doc, 
        "Tiến hành chạy thử nghiệm stress-test hệ thống hoạt động liên tục trong 24 giờ. Thực hiện đo đạc dòng điện tiêu thụ ở các chế độ làm việc và kiểm tra nhiệt độ hoạt động của chip vi xử lý ESP32-S3 bằng súng đo nhiệt hồng ngoại để đánh giá độ ổn định và an toàn nhiệt.",
        bold_prefix="• Đo đạc nhiệt độ và điện năng tiêu thụ (Stress Test): "
    )
    
    # --- SECTION 2 ---
    add_heading_styled(doc, "2. Phân công thực hiện", level=1)
    
    add_heading_styled(doc, "Trần Văn Chót:", level=2)
    add_paragraph_styled(doc, "Lập trình Flask Server trung tâm, cấu hình cơ sở dữ liệu SQLite faces.db và các API Webhook xử lý dữ liệu.", bullet=True)
    add_paragraph_styled(doc, "Phát triển bộ chuyển tiếp log từ xa qua giao thức UDP (cổng 8091) trên ESP32-S3 bằng cách cấu hình ghi đè hàm vprintf.", bullet=True)
    add_paragraph_styled(doc, "Thiết kế mạch đo đạc dòng điện tiêu thụ của toàn bộ hệ thống bằng Ampe kế ở các chế độ hoạt động khác nhau.", bullet=True)
    
    add_heading_styled(doc, "Dương Thị Quỳnh Châm:", level=2)
    add_paragraph_styled(doc, "Lập trình giao diện Web Dashboard index.html hiển thị luồng stream video trực tuyến, bảng log điểm danh và các nút cấu hình đổi tên.", bullet=True)
    add_paragraph_styled(doc, "Cấu hình chức năng nén và gửi ảnh JPEG (Multipart HTTP POST) từ thiết bị biên lên máy chủ Flask sau khi đăng ký thành công.", bullet=True)
    add_paragraph_styled(doc, "Thực hiện quá trình stress-test 24 giờ liên tục, ghi nhận nhiệt độ chip và theo dõi hiện tượng rò rỉ bộ nhớ (Memory Leak).", bullet=True)
    
    # --- SECTION 3 ---
    add_heading_styled(doc, "3. Báo cáo quá trình thực hiện", level=1)
    
    add_heading_styled(doc, "3.1. Xây dựng máy chủ Flask trung tâm và Web Dashboard:", level=2)
    add_paragraph_styled(doc, 
        "Máy chủ Flask được lập trình bằng Python chạy độc lập trên máy chủ trung tâm (PC/Mac) kết nối cùng mạng Wi-Fi cục bộ với thiết bị nhúng. Cơ sở dữ liệu SQLite faces.db lưu trữ bảng faces với các cột ID và Name. Web Dashboard cung cấp các tính năng quản lý trực quan: hiển thị danh sách thành viên cùng ảnh chân dung gốc tương ứng được lưu trữ trong thư mục static/images, bảng nhật ký điểm danh thời gian thực, và hỗ trợ chỉnh sửa đổi tên trực tiếp trên Web. Lệnh đổi tên (Rename) trên Web sẽ được gửi đồng bộ ngược lại bộ nhớ Flash NVS của thiết bị nhúng qua API HTTP POST '/api/faces/rename'."
    )
    
    # Image 1
    add_image_styled(doc, os.path.join(pic_dir, "IMG_2817.jpeg"), "Hình 1: Giao diện Web Dashboard quản lý cơ sở dữ liệu khuôn mặt và giám sát nhật ký trực tuyến")
    
    add_heading_styled(doc, "3.2. Triển khai giao thức tự động kết nối mạng UDP Auto-Discovery:", level=2)
    add_paragraph_styled(doc, 
        "• Nguyên lý hoạt động: Để thiết bị hoạt động linh hoạt mà không cần cấu hình cứng địa chỉ IP của server trong code, nhóm đã triển khai giao thức UDP Auto-Discovery. Khi thiết bị biên ESP32-S3 khởi động và kết nối Wi-Fi thành công, nó sẽ tự động tạo một FreeRTOS Task chạy ngầm gửi gói tin Broadcast UDP chứa thông điệp 'ESP32_DISCOVER:<IP_thiết_bị>' qua cổng 8090 đến địa chỉ quảng bá 255.255.255.255 định kỳ mỗi 3 giây."
    )
    add_paragraph_styled(doc, 
        "• Phản hồi từ Server: Flask Server bind cổng 8090 để lắng nghe gói tin broadcast này. Khi nhận được thông điệp từ thiết bị biên, Server sẽ cập nhật IP của ESP32-S3 và gửi phản hồi lại bằng gói tin Unicast UDP chứa thông điệp 'SERVER_IP:<IP_máy_chủ>'."
    )
    add_paragraph_styled(doc, 
        "• Kết quả: ESP32-S3 nhận được phản hồi sẽ lưu lại IP máy chủ, thực hiện gửi Webhook khởi động (startup webhook) để xác nhận kết nối, đồng thời tự động đóng socket và giải phóng task FreeRTOS để tiết kiệm 16KB bộ nhớ RAM."
    )
    
    add_heading_styled(doc, "3.3. Cơ chế Đăng ký nhanh và Đồng bộ hóa khuôn mặt:", level=2)
    add_paragraph_styled(doc, 
        "• Đăng ký người dùng mới trực tiếp (Local Enroll): Khi nhấn đúp nút BOOT (GPIO 0) trên thiết bị nhúng hoặc chọn kích hoạt trên Web Dashboard, hệ thống nhúng sẽ chuyển sang chế độ Enroll. Camera chụp ảnh chân dung, bộ phát hiện tìm khuôn mặt, bộ trích xuất đặc trưng MobileNetV2 tính toán vector 128 chiều lượng hóa INT8 lưu vào LittleFS trên Flash, đồng thời NVS ghi nhận ánh xạ ID - Name. Ngay sau đó, thiết bị nhúng tiến hành chuyển đổi frame ảnh sang định dạng JPEG và gửi HTTP POST truyền tải tệp nhị phân ảnh chân dung gốc lên máy chủ Flask để lưu trữ phục vụ việc hiển thị trên Web."
    )
    add_paragraph_styled(doc, 
        "• Chuyển tiếp log gỡ lỗi từ xa qua UDP (Remote Logging): Nhóm đã cấu hình ghi đè bộ xuất log của ESP-IDF thông qua hàm esp_log_set_vprintf. Toàn bộ thông tin log debug của thiết bị nhúng biên được chuyển tiếp bất đồng bộ qua cổng UDP 8091 về máy chủ Flask và hiển thị trực tiếp lên Web Dashboard, giúp giám sát và gỡ lỗi hệ thống từ xa mà không cần cắm cáp nạp USB."
    )
    
    add_heading_styled(doc, "3.4. Kết quả đo đạc nhiệt độ và điện năng tiêu thụ (Stress Test):", level=2)
    add_paragraph_styled(doc, 
        "• Đo đạc dòng điện tiêu thụ (nguồn cấp 5V): Nhóm tiến hành đo đạc dòng điện tiêu thụ của mạch nhúng bằng Ampe kế ở các chế độ hoạt động: (1) Chế độ chờ Idle (WiFi kết nối, màn hình LCD bật sáng, camera hoạt động): dòng điện dao động ổn định trong khoảng 190mA - 220mA (tương đương công suất tiêu thụ ~1.0W); (2) Chế độ chạy mô hình AI nhận diện (MobileNetV2 chạy liên tục trên Core 1): dòng điện tăng vọt và đạt đỉnh ở mức 360mA - 410mA (tương đương công suất ~2.0W); (3) Khi kích hoạt Relay mở khóa: dòng điện tăng thêm 70mA."
    )
    add_paragraph_styled(doc, 
        "• Đo đạc nhiệt độ chip xử lý (nhiệt độ phòng 28°C): Sử dụng thiết bị đo nhiệt hồng ngoại đo trực tiếp trên bề mặt vỏ chip ESP32-S3-WROOM-1: Sau 10 phút hoạt động, nhiệt độ chip duy trì ở mức 39°C. Khi hệ thống nhận diện khuôn mặt liên tục trong 1 giờ, nhiệt độ bề mặt chip tăng lên và ổn định ở khoảng 46°C - 49°C. Nhóm đã thực hiện stress-test hệ thống hoạt động liên tục trong 24 giờ, nhiệt độ tối đa ghi nhận là 51.2°C, hoàn toàn nằm trong giới hạn an toàn nhiệt của linh kiện bán dẫn nhúng (< 85°C). Giám sát heap RAM cho thấy bộ nhớ khả dụng duy trì ở mức ổn định ~4.2MB RAM, không xảy ra hiện tượng rò rỉ bộ nhớ (Memory Leak) khi chạy đa nhiệm FreeRTOS trong thời gian dài."
    )
    
    # Image 2
    add_image_styled(doc, os.path.join(pic_dir, "IMG_2820.jpeg"), "Hình 2: Thử nghiệm đo đạc điện năng tiêu thụ và nhiệt độ hoạt động của thiết bị biên")
    
    # --- SECTION 4 ---
    add_heading_styled(doc, "4. Dự kiến công việc tuần tiếp theo", level=1)
    
    add_paragraph_styled(doc, "Tiến hành đo đạc và đánh giá chi tiết độ chính xác của hệ thống nhận diện khuôn mặt thời gian thực trong các điều kiện môi trường ánh sáng phức tạp khác nhau (ngược sáng, bóng tối, ánh sáng đèn huỳnh quang). Thu thập số liệu thống kê.", bullet=True)
    add_paragraph_styled(doc, "Tổng hợp hình ảnh thực tế, vẽ các biểu đồ hiệu năng, hoàn thiện toàn bộ nội dung báo cáo tổng kết đồ án môn học và thiết kế slide báo cáo.", bullet=True)
    add_paragraph_styled(doc, "Hoàn thiện đóng gói phần cứng sản phẩm demo cơ khí hoàn chỉnh (cố định mạch điện, cố định các cổng kết nối ngoại vi) để chuẩn bị cho buổi báo cáo nghiệm thu đồ án trước hội đồng.", bullet=True)
    
    print(f"Saving report to {docx_dst}...")
    doc.save(docx_dst)
    print("Report generated successfully!")

if __name__ == "__main__":
    main()
